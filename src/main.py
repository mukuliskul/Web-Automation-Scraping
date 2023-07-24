from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import sys
import time
from docx import Document
import xlwings as xw
import math
import selenium

excel1='2016-17 green.xlsx'
excel2='depts 51-74 year 18-19.xlsx'
excel3='2018-19 depts 61-70 PU Publications.xlsx'

wb = xw.Book(excel3)
sht = wb.sheets('Bibliogrphic details')

deptname='zoology'
difference=int(387)

filecount=int(input('Enter FILE NUMBER --> ' )) #check the numbering once of the continous word files
cellcount=filecount+difference
print(' ')

title_list = sht.range(f'C{cellcount}:C1000').value
""" title_list=["Role of GSK 3â in regulation of canonical Wnt/â-Catenin signaling and PI 3K/AKToncognenic pathway in colon cancer. Cancer Investigation.2017, DOI 1080/07357907. 2017.1337783.","Role of angeogenic factors of herbal origin in regulation of molecular pathwaysthat control tumor angiogenesis Tumor Biology. 37, 14341-354."] """

driver = webdriver.Chrome('/Users/mukulsharma/Desktop/CS/python/automation/chromedriver')
wait = WebDriverWait(driver, 200)
driver.get('https://scholar.google.com')
time.sleep(0.8)
cellnumbercount=int(cellcount)

for title in title_list:
    cellnumbercount+=1
    renamed=title[0:30]
    APAdocx=f'{filecount} {deptname} {renamed}.docx'
    APAlist=[]
    count=0

    searchBox = wait.until(EC.element_to_be_clickable((By.XPATH, '//*[@id="gs_hdr_tsi"]')))    
    searchBox.send_keys(title)
    time.sleep(1)
    search= webdriver.ActionChains(driver).send_keys(Keys.ENTER).perform()
    command=input("Enter Operation --> ")
    if command=='n' or command=='N':
        filecount+=1
        driver.get('https://scholar.google.com')
        print(cellnumbercount)
        print(filecount)
        continue
        

    else:
        citedbyLink = wait.until(EC.element_to_be_clickable((By.PARTIAL_LINK_TEXT, 'Cited by')))
        pagetotaltxt=citedbyLink.text
        pagetotal=pagetotaltxt.lstrip('Cited by ')
        citedbyLink.click()
    
        citeLINK=driver.find_elements(By.LINK_TEXT, 'Cite')
        for i in citeLINK:
            count+=1
            time.sleep(1)
            i.click()
            time.sleep(1)
            APAclick=wait.until(EC.element_to_be_clickable((By.XPATH, '//*[@id="gs_citt"]/table/tbody/tr[2]/td/div')))
            APAtext=str(count)+'. '+str(APAclick.text)
            APAlist.append(APAtext)
            time.sleep(1)
            webdriver.ActionChains(driver).send_keys(Keys.ESCAPE).perform()
    
        totaldecimal=int(pagetotal)/10
        totalroundoff=math.ceil(float(totaldecimal))
        finaltotal=int(totalroundoff)+1
        finaltotal2=int(totalroundoff)-1
        
        if finaltotal<=11:
            for i in range(2,int(finaltotal)):
                next=driver.find_element(By.LINK_TEXT, str(i)).click()
                citeLINK=[]
                citeLINK=driver.find_elements(By.LINK_TEXT, 'Cite')
                for i in citeLINK:
                    count+=1
                    time.sleep(1.5)
                    i.click()
                    time.sleep(1)
                    APAclick=wait.until(EC.element_to_be_clickable((By.XPATH, '//*[@id="gs_citt"]/table/tbody/tr[2]/td/div')))
                    #APAclick=driver.find_element(By.XPATH, '//*[@id="gs_citt"]/table/tbody/tr[2]/td/div')
                    APAtext=str(count)+'. '+str(APAclick.text)
                    APAlist.append(APAtext)
                    time.sleep(1)
                    webdriver.ActionChains(driver).send_keys(Keys.ESCAPE).perform()
        else:
            for i in range(int(finaltotal2)):
                next=driver.find_element(By.LINK_TEXT, 'Next').click()
                citeLINK=[]
                citeLINK=driver.find_elements(By.LINK_TEXT, 'Cite')
                for i in citeLINK:
                    count+=1
                    time.sleep(1.5)
                    i.click()
                    time.sleep(1)
                    APAclick=wait.until(EC.element_to_be_clickable((By.XPATH, '//*[@id="gs_citt"]/table/tbody/tr[2]/td/div')))
                    #APAclick=driver.find_element(By.XPATH, '//*[@id="gs_citt"]/table/tbody/tr[2]/td/div')
                    APAtext=str(count)+'. '+str(APAclick.text)
                    APAlist.append(APAtext)
                    time.sleep(1)
                    webdriver.ActionChains(driver).send_keys(Keys.ESCAPE).perform()
                
    
    
        document=Document()
        document.add_paragraph(title)
        for i in APAlist:
            para=document.add_paragraph(i)
        document.save(APAdocx)
        
        filecount+=1
        driver.get('https://scholar.google.com')
        
        print(cellnumbercount)
        print(filecount)





